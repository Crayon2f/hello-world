# coding=utf-8
import shutil
import traceback

import pymysql
import sys
import os
import math
import time
from kit import config_kit
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from oss import oss_kit
from import_zy import get_region
import importlib

importlib.reload(sys)

connection = pymysql.connect(host='mt-58art-database-open.mysql.rds.aliyuncs.com',
                             port=3306,
                             user='mt_art58',
                             passwd='Admin_58art',
                             db='art58',
                             charset='utf8')
cursor = connection.cursor(pymysql.cursors.DictCursor)

# artwork_dir = '/Volumes/Crayon2f/artworks/'
artwork_dir = u'e:\\artworks\\'
final_artwork_dir = u'e:\\final_round\\'
change_sign = '-'
activity_id = config_kit.CONFIG.get('activity', 'id')


def create_dir():
    """
    创建所有参展的艺术家文件夹
    :return:
    """
    sql = "SELECT user_id FROM activity_registration WHERE is_virtual = 0  and activity_id = %s AND user_id <> ''" % activity_id
    cursor.execute(sql)
    user_list = cursor.fetchall()
    if len(user_list) > 0:
        if not os.path.exists(artwork_dir):
            os.mkdir(artwork_dir)
    for user in user_list:
        os.mkdir(artwork_dir + user['user_id'])


def get_artwork():
    """
    获取当前海选所有作品
    :return:
    """
    sql = "SELECT  `NAME` AS name, ID id, AUTHOR author, `client` FROM voting_artwork WHERE ACTIVITY_ID = %s " % activity_id
    cursor.execute(sql)
    return cursor.fetchall()


def download_artwork_form_oss():
    """
    下载OSS的作品
    :return:
    """
    artwork_list = get_artwork()
    index_out = 0
    bucket_name = 'mt-original'
    _oss_kit = oss_kit.OssKit(bucket_name)
    for artwork in artwork_list:
        index_out += 1
        print(index_out)
        dir_name = artwork_dir + artwork['author']
        artwork_name = translate_sign(artwork['name'].strip())
        artwork_path = '%s/%s.jpg' % (dir_name, artwork_name)
        if os.path.exists(artwork_path):
            continue
            # artwork_path = recursion(artwork_path, 0)
        key_list = get_oss_key(artwork['id'], bucket_name)
        try:
            if len(key_list) > 0:
                key = key_list[0]
                if not os.path.exists(artwork_path):
                    _oss_kit.download(str(key['key']), artwork_path)
                if len(key_list) > 1:
                    for index, append_key in enumerate(key_list[1:]):
                        append_dir = '%s/detail/' % dir_name
                        if not os.path.exists(append_dir):
                            os.mkdir(append_dir)
                        append_path = '%s/%s[%d].jpg' % (append_dir, artwork_name, index + 1)
                        if not os.path.exists(append_path):
                            _oss_kit.download(str(append_key['key']), append_path)
            else:
                print(artwork['id'] + ' has not oss_key')
        except BaseException as exception:
            print(exception)
            print('===== %s =====' % artwork['id'])
            traceback.print_exc()


def get_oss_key(artwork_id, bucket_name):
    """
    获取OSS地址
    :param artwork_id: 作品ID
    :param bucket_name: bucket
    :return:
    """
    sql = "select key_value as `key` from art_images where source_id = '%s' and " \
          "bucket_name = '%s' and RULE_CODE = 'original' order by category desc " % (artwork_id, bucket_name)
    cursor.execute(sql)
    return cursor.fetchall()


def translate_sign(origin):
    """
    转换特殊符号
    :param origin:
    :return:
    """
    if origin == '':
        return origin
    return origin.replace('<', change_sign).replace('>', change_sign).replace(':', change_sign) \
        .replace('"', change_sign).replace('/', change_sign).replace('\\', change_sign) \
        .replace('|', change_sign).replace('?', change_sign).replace('*', change_sign)


def export_resume():
    """
    导出简历
    :return:
    """
    region_dict = get_region.get_name()
    sql = "SELECT ar.user_id, ar.real_name, ar.birthday, ar.live_place AS live_place," \
          " ar.terminal, su.BIRTHPLACE as birth_place FROM " \
          "activity_registration ar LEFT JOIN sys_user su ON ar.user_id = su.ID " \
          "WHERE ar.activity_id = %s and ar.user_id <> '' and ar.user_id is not null" % activity_id
    cursor.execute(sql)
    user_list = cursor.fetchall()
    index = 0
    for user in user_list:
        index += 1
        print(index)
        personal_exhibition = []
        joint_exhibition = []
        try:
            if user['terminal'] != 'zai_yi':
                show_sql = "select NAME as name, SPACENAME as space_name, CITY as city, `YEAR` as  `year`, " \
                           "TYPE as type from art_show_extern where CRT_TELLER_ID = '%s' ORDER BY YEAR desc" % \
                           user['user_id']
                cursor.execute(show_sql)
                show_list = cursor.fetchall()
                if len(show_list) > 0:
                    for show in show_list:
                        if int(show['type']) == 0:
                            personal_exhibition.append(show)
                        else:
                            joint_exhibition.append(show)
            region_id = user['live_place'] if user['live_place'] and user['live_place'] != '' else None
            region_name = u'未知' if (None is region_id) or (type(region_id) is float and math.isnan(region_id)) else \
                region_dict[region_id]
            real_name = str(user['real_name']).strip()
            out_path = u"%s%s\\个人简历.docx" % (artwork_dir, user['user_id'])
            if not os.path.exists(artwork_dir + user['user_id']):
                out_path = u"%s%s.docx" % (artwork_dir, translate_sign(real_name))
            education_list = get_education_list(user['user_id'])
            if education_list and len(education_list) > 0:
                write_word(joint_exhibition, personal_exhibition, out_path, real_name, user['birthday'], region_name,
                           user['birth_place'], education_list)
            print(user['real_name'] + ' resume export successful ')
        except BaseException as exception:
            print("export word error : ".format(exception))
            traceback.print_exc()
            print("export word error user_id ==> %s" % user['user_id'])


def get_education_list(user_id):
    """
    获取学历集合
    :param user_id: 用户ID
    :return:
    """
    sql = "select * from artist_education where user_id = '%s' order by education asc" % user_id
    cursor.execute(sql)
    return cursor.fetchall()


def write_word(joint_exhibition, personal_exhibition, out_path, real_name='', birthday='', live_place='',
               birth_place='', education_list=None):
    """
    生成 word
    :param joint_exhibition: 联展
    :param personal_exhibition: 个展
    :param out_path: 输出目录
    :param real_name: 真是姓名
    :param birthday: 生日
    :param live_place: 现居住地
    :param birth_place: 出生地
    :param education_list: 学历集合
    :return:
    """
    if education_list is None:
        education_list = []
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    paragraph = document.add_paragraph()
    if real_name is None:
        real_name = ''
    name = paragraph.add_run(u'%s' % real_name)
    name.font.size = Pt(12)
    name.bold = True

    paragraph = document.add_paragraph()
    paragraph.add_run(u'%s生于%s' % (birthday, birth_place))
    paragraph = document.add_paragraph()
    if education_list and len(education_list) > 0:
        for education in education_list:
            education_str = ''
            if education['education'] == 2000:
                education_str = u'学士'
            elif education['education'] == 3000:
                education_str = u'硕士'
            elif education['education'] == 4000:
                education_str = u'博士'
            if education_str != '':
                if education['years'] > time.localtime().tm_year:
                    paragraph.add_run(u'%d将毕业于%s%s系 获%s学位' % (
                        education['years'], education['university'], education['series'], education_str))
                else:
                    paragraph.add_run(u'%d毕业于%s%s系 获%s学位' % (
                        education['years'], education['university'], education['series'], education_str))
                paragraph = document.add_paragraph()
    paragraph.add_run(u'现工作生活于%s' % live_place)

    paragraph = document.add_paragraph('\n')
    name = paragraph.add_run(u'个展')
    name.font.size = Pt(12)
    name.bold = True

    generate_show(document, personal_exhibition)

    paragraph = document.add_paragraph('\n')
    name = paragraph.add_run(u'联展')
    name.font.size = Pt(12)
    name.bold = True

    generate_show(document, joint_exhibition)
    document.save(out_path)


def generate_show(document, exhibition_list):
    """
    生成展览
    :param paragraph: 换行符
    :param document: 当前 word
    :param exhibition_list: 展览
    :return:
    """
    exhibition_map = {}
    exhibition_list = list(filter(lambda e: e['name'] != '无', exhibition_list))
    if exhibition_list and len(exhibition_list) > 0:
        year_list = []
        for exhibition in exhibition_list:
            year = exhibition['year']
            if year not in exhibition_map:
                exhibition_map[year] = [exhibition]
            else:
                exhibition_map[year].append(exhibition)
            if year not in year_list:
                year_list.append(year)
        for year in year_list:
            paragraph = document.add_paragraph()
            paragraph.add_run(u'%s' % year)
            for exhibition in exhibition_map[year]:
                paragraph = document.add_paragraph()
                paragraph.add_run(u'%s，%s，%s' % (exhibition['name'], exhibition['space_name'], exhibition['city']))
    else:
        paragraph = document.add_paragraph()
        paragraph.add_run(u'无')


def recursion(path, index):
    """
    递归找重复的名字
    :param path: 路径
    :param index: 当前重复的次数
    :return:
    """
    if os.path.exists(path):
        index += 1
        path = '%s(%d)' % (path, index)
        return recursion(path, index)
    else:
        return path


def translate():
    """
    将文件夹ID转换为汉字
    :return:
    """
    sql = "SELECT user_id, real_name FROM activity_registration WHERE is_virtual = 0 and  activity_id = %s and user_id <> '';" % activity_id
    cursor.execute(sql)
    user_list = cursor.fetchall()
    for user in user_list:
        author_dir = os.path.join(artwork_dir, user['user_id'])
        # author_dir = os.path.join(final_artwork_dir, user['user_id'])
        if os.path.exists(author_dir):
            new_author_dir = os.path.join(artwork_dir, translate_sign(user['real_name'].strip()))
            # new_author_dir = os.path.join(final_artwork_dir, translate_sign(user['real_name'].strip()))
            if os.path.exists(new_author_dir):
                new_author_dir = recursion(new_author_dir, 0)
            print(new_author_dir)
            try:
                os.rename(author_dir, new_author_dir)
                print(user['real_name'] + ' success ')
            except BaseException as e:
                print(e)
                traceback.print_exc()
                print(user['user_id'] + ' fail')


def export_attachment():
    """
    导出附件
    :return:
    """
    sql = 'select user_id, bucket_name, key_value from activity_attachment where activity_id = %s' % activity_id
    cursor.execute(sql)
    attachment_list = cursor.fetchall()
    for attachment in attachment_list:
        author_dir = os.path.join(artwork_dir, attachment['user_id'])
        oss_kit_ = oss_kit.OssKit(attachment['bucket_name'])
        key = str(attachment['key_value'])
        oss_kit_.download(key, os.path.join(author_dir, u'方案' + key[key.index('.'):]))
    return


def fix_name_duplicate_artwork():
    """
    修改名字重复的作品
    """
    dir_list = os.listdir(artwork_dir)
    for user_dir in dir_list:
        artwork_list = os.listdir(os.path.join(artwork_dir, user_dir))
        artwork_map = {}
        for artwork in artwork_list:
            artwork_path = os.path.join(artwork_dir, user_dir, artwork)
            artwork = artwork[0:artwork.find('(1)')]
            if artwork not in artwork_map:
                artwork_map[artwork] = [artwork_path]
            else:
                artwork_map[artwork].append(artwork_path)
            # if os.path.isfile(artwork_path):
            #     if artwork_path.find('(1)') > 0:
            #         print(artwork_path
            #
            # suffix = os.path.splitext(artwork_path)[1]
            # if suffix != '.jpg':
            #     index = artwork_path.find('.jpg')
            #     new_path = artwork_path[0:index] + artwork_path[index + 4:] + '.jpg'
            #     print(artwork_path + " ==> " + new_path
            #
            #     os.rename(artwork_path, new_path)
        for key in artwork_map:
            duplicate_artwork = artwork_map[key]
            if len(duplicate_artwork) > 1:
                i = 1
                for name in duplicate_artwork:
                    new_path = os.path.join(name[0:name.find('(1)')] + '[' + str(i) + ']') + '.jpg'
                    print(name + " ==> " + new_path)
                    os.rename(name, new_path)
                    i += 1


def copy_designation_round_artwork():
    """
    从所有的艺术家作品中 复制指定轮次的作品
    """
    sql = "select user_id from voting where show_id = %s and round > 4 and voting_user_id = 60" % activity_id
    cursor.execute(sql)
    user_list = cursor.fetchall()
    for user in user_list:
        target_dir = os.path.join(final_artwork_dir, user['user_id'])
        source_dir = os.path.join(artwork_dir, user['user_id'])
        shutil.copytree(source_dir, target_dir)


if __name__ == '__main__':
    # 创建文件夹uuid形式
    # create_dir()
    # cs阿里云下载图片
    # download_artwork_form_oss()
    # 处理名称重复的图片
    # fix_name_duplicate_artwork()
    # export_resume()
    # export_attachment()
    translate()
    # copy_designation_round_artwork()
    # dir_list = os.listdir(os.path.join(final_artwork_dir))
    # i = 1
    # for user_dir in dir_list:
    #     print(i)
    #     duplicate_path = os.path.join(artwork_dir, user_dir + "(1)")
    #     if os.path.exists(duplicate_path):
    #         shutil.rmtree(duplicate_path)
    #     i += 1
