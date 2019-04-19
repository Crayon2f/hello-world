# coding=utf-8
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
name = paragraph.add_run(u'宋继瑞')
name.font.size = Pt(12)
name.bold = True

paragraph = document.add_paragraph()
paragraph.add_run(u'1986生于山东青州')
paragraph = document.add_paragraph()
paragraph.add_run(u'2006毕业于中央美术学院壁画系 获学士学位')
paragraph = document.add_paragraph()
paragraph.add_run(u'2010毕业于中央美术学院壁画系 获硕士学位')
paragraph = document.add_paragraph()
paragraph.add_run(u'2015毕业于中央美术学院壁画系 获博士学位')
paragraph = document.add_paragraph()
paragraph.add_run(u'现工作生活于北京')

paragraph = document.add_paragraph('\n')
name = paragraph.add_run(u'个展')
name.font.size = Pt(12)
name.bold = True

paragraph = document.add_paragraph()
paragraph.add_run(u'2015')
paragraph = document.add_paragraph()
paragraph.add_run(u'两种关系——迟群个人作品展，名泰空间，北京')
paragraph = document.add_paragraph()
paragraph.add_run(u'青春礼赞——中韩青年艺术展，弘益大学现代美术馆，首尔，韩国')
paragraph = document.add_paragraph()
paragraph.add_run(u'两种关系——迟群个人作品展，名泰空间，北京')

paragraph = document.add_paragraph('\n')
name = paragraph.add_run(u'群展')
name.font.size = Pt(12)
name.font.name = u'宋体'
name.bold = True

paragraph = document.add_paragraph()
paragraph.add_run(u'2015')
paragraph = document.add_paragraph()
paragraph.add_run(u'青春礼赞——中韩青年艺术展，弘益大学现代美术馆，首尔，韩国')
paragraph = document.add_paragraph()
paragraph.add_run(u'两种关系——迟群个人作品展，名泰空间，北京')
paragraph = document.add_paragraph()
paragraph.add_run(u'青春礼赞——中韩青年艺术展，弘益大学现代美术馆，首尔，韩国')

document.save("C:\\Users\\Lenovo\\Desktop\\test3.docx")
